"""Iteration 78 — New Contract pricing table + amount-in-words + no yellow highlight.

Tests:
 1. Pricing table row 1 dynamic cells (Stend №, Xidmət, Ölçü vahidi, Qiymət, ƏDV, Yekun)
 2. 'Məbləğ yazı ilə: …' paragraph reflects the computed total in Azerbaijani words
 3. Two-phase token substitution (price containing '50' substring not corrupted by stand_no rule)
 4. All yellow highlights removed from generated DOCX
 5. No PUA tokens (U+E000..U+E01F) leaked
 6. Regression — Yeni Əlavə Müqavilə (addendum) still generates valid DOCX
"""
import io
import os
import re
import pytest
import requests
from docx import Document

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL").rstrip("/")
API = f"{BASE_URL}/api"

PUA_RE = re.compile(r"[\uE000-\uE01F]")


@pytest.fixture(scope="module")
def client():
    r = requests.post(f"{API}/auth/login",
                      json={"email": "settings@marsol.az", "password": "marsol123"},
                      timeout=15)
    assert r.status_code == 200, r.text
    tok = r.json()["access_token"]
    s = requests.Session()
    s.headers.update({"Authorization": f"Bearer {tok}",
                      "Content-Type": "application/json"})
    return s


@pytest.fixture(scope="module")
def created_ids(client):
    ids = []
    yield ids
    for cid in ids:
        try:
            client.delete(f"{API}/contracts/{cid}")
        except Exception:
            pass


def _payload(**over):
    p = {
        "contract_number": "TS_QA/27",
        "contract_date": "2027-02-15",
        "sifarisci_company": "TEST_QA MMC",
        "sifarisci_voen": "1234567890",
        "sifarisci_authorized": "TEST Executor",
        "iban": "AZ00TEST",
        "bank_name": "TEST Bank",
        "branch_code": "888",
        "bank_voen": "7777777777",
        "correspondent_account": "AZ00CORR",
        "swift": "TESTAZ2X",
        "stand_number": "7",
        "stand_width": 3,
        "stand_length": 4,
        "price": 15000,
        "vat_enabled": True,
        "vat_rate": 18,
    }
    p.update(over)
    return p


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
                for tt in cell.tables:  # nested
                    for rr in tt.rows:
                        for cc in rr.cells:
                            parts.append(cc.text)
    return "\n".join(parts)


def _count_highlights(doc):
    count = 0
    def walk_paras(paras):
        nonlocal count
        for p in paras:
            for run in p.runs:
                if run.font.highlight_color is not None:
                    count += 1
    walk_paras(doc.paragraphs)
    def walk_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    walk_paras(cell.paragraphs)
                    walk_tables(cell.tables)
    walk_tables(doc.tables)
    return count


class TestPricingAndFormat:

    def test_pricing_15000(self, client, created_ids):
        r = client.post(f"{API}/contracts/new", json=_payload())
        assert r.status_code == 200, r.text
        cid = r.json()["id"]
        created_ids.append(cid)

        d = client.get(f"{API}/contracts/{cid}/download")
        assert d.status_code == 200
        doc = Document(io.BytesIO(d.content))

        # Table 1 (index 1) row 1 pricing values
        pricing_tbl = doc.tables[1]
        row1 = pricing_tbl.rows[1]
        cells = [c.text.strip() for c in row1.cells]
        print("Pricing row cells:", cells)

        assert cells[0] == "7", f"Stend № expected '7', got {cells[0]!r}"
        assert "Sərgidə iştirak" in cells[1], f"Xidmət cell: {cells[1]!r}"
        assert cells[2] in ("x/h", "ə/v"), f"Ölçü vahidi cell: {cells[2]!r}"
        assert cells[3] == "15 000.00", f"Qiymət expected '15 000.00', got {cells[3]!r}"
        assert cells[4] == "2 700.00", f"ƏDV expected '2 700.00', got {cells[4]!r}"
        assert cells[5] == "17 700.00", f"Yekun expected '17 700.00', got {cells[5]!r}"

        # amount in words paragraph
        text = _all_text(doc)
        expected_words = "on yeddi min yeddi yüz manat 00 qəpik"
        assert expected_words in text, (
            f"Expected 'Məbləğ yazı ilə: {expected_words}' not found. "
            f"Search 'Məbləğ' snippet: "
            f"{[ln for ln in text.split(chr(10)) if 'Məbləğ' in ln or 'manat' in ln][:5]}"
        )

        # no PUA tokens leaked
        m = PUA_RE.search(text)
        assert m is None, f"PUA token leaked: {m.group(0)!r} at pos {m.start()}"

        # no yellow highlight
        hl = _count_highlights(doc)
        assert hl == 0, f"Expected 0 highlighted runs, got {hl}"

        # No leftover template static values
        for leftover in ["90 000.00", "16 200.00", "106 200.00",
                         "yüz altı min iki yüz manat"]:
            assert leftover not in text, f"Leftover template value {leftover!r} present"

    def test_pricing_5500_two_phase(self, client, created_ids):
        """price=5500 contains '50' substring — must not get corrupted by stand_no rule."""
        r = client.post(f"{API}/contracts/new", json=_payload(
            contract_number="TS_QA2/27",
            price=5500,
            stand_number="12",
        ))
        assert r.status_code == 200, r.text
        cid = r.json()["id"]
        created_ids.append(cid)

        d = client.get(f"{API}/contracts/{cid}/download")
        assert d.status_code == 200
        doc = Document(io.BytesIO(d.content))

        row1 = doc.tables[1].rows[1]
        cells = [c.text.strip() for c in row1.cells]
        print("5500 pricing cells:", cells)

        assert cells[0] == "12", f"Stend № expected '12', got {cells[0]!r}"
        assert cells[3] == "5 500.00", f"Qiymət expected '5 500.00', got {cells[3]!r}"
        assert cells[4] == "990.00", f"ƏDV expected '990.00', got {cells[4]!r}"
        assert cells[5] == "6 490.00", f"Yekun expected '6 490.00', got {cells[5]!r}"

        text = _all_text(doc)
        # 6490 = altı min dörd yüz doxsan
        assert "altı min dörd yüz doxsan manat 00 qəpik" in text, (
            "Expected words 'altı min dörd yüz doxsan manat 00 qəpik' missing. "
            f"Sample lines: {[ln for ln in text.split(chr(10)) if 'manat' in ln][:5]}"
        )
        assert PUA_RE.search(text) is None, "PUA token leaked"
        assert _count_highlights(doc) == 0, "Yellow highlight not removed"

    def test_addendum_regression(self, client, created_ids):
        r = client.post(f"{API}/contracts/addendum", json={
            "parent_contract_number": "TS_QA/27",
            "parent_contract_date": "2027-02-15",
            "addendum_date": "2027-03-01",
            "sifarisci_company": "TEST_QA MMC",
            "sifarisci_voen": "1234567890",
            "sifarisci_authorized": "TEST Executor",
            "stand_number": "7",
            "stand_width": 3,
            "stand_length": 4,
            "exhibition_name": "Yerli sərgi",
            "exhibition_start": "2027-06-23",
            "exhibition_end": "2027-06-26",
            "pricing": {"price_net": 15000, "vat_enabled": True, "vat_rate": 18},
        })
        assert r.status_code == 200, r.text
        cid = r.json()["id"]
        created_ids.append(cid)

        d = client.get(f"{API}/contracts/{cid}/download")
        assert d.status_code == 200
        assert len(d.content) > 5000
        doc = Document(io.BytesIO(d.content))
        text = _all_text(doc)
        # Just ensure it opens & no PUA leaks
        assert PUA_RE.search(text) is None, "PUA token leaked in addendum"
