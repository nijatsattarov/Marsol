"""Iteration 77 — New Contract (Yeni Müqavilə) endpoint tests."""
import io
import os
import pytest
import requests
from docx import Document
from openpyxl import load_workbook

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL").rstrip("/")
API = f"{BASE_URL}/api"


@pytest.fixture(scope="module")
def token():
    r = requests.post(f"{API}/auth/login",
                      json={"email": "settings@marsol.az", "password": "marsol123"},
                      timeout=15)
    assert r.status_code == 200, r.text
    return r.json()["access_token"]


@pytest.fixture(scope="module")
def client(token):
    s = requests.Session()
    s.headers.update({"Authorization": f"Bearer {token}",
                      "Content-Type": "application/json"})
    return s


@pytest.fixture(scope="module")
def created_ids():
    ids = []
    yield ids
    # teardown - cleanup TEST_ contracts
    s = requests.Session()
    r = s.post(f"{API}/auth/login",
               json={"email": "settings@marsol.az", "password": "marsol123"})
    tok = r.json().get("access_token")
    s.headers.update({"Authorization": f"Bearer {tok}"})
    for cid in ids:
        s.delete(f"{API}/contracts/{cid}")


def _payload(**over):
    p = {
        "contract_number": "TS_QA/27",
        "contract_date": "2027-02-15",
        "sifarisci_company": "TEST_QA MMC",
        "sifarisci_voen": "9999999999",
        "sifarisci_authorized": "TEST Executor",
        "iban": "AZ00TEST0000000000000000TEST",
        "bank_name": "TEST Bank",
        "branch_code": "888",
        "bank_voen": "7777777777",
        "correspondent_account": "AZ00TEST0000CORR",
        "swift": "TESTAZ2X",
        "stand_number": "S-99",
        "stand_width": 4,
        "stand_length": 3,
        "price": 5000,
        "vat_enabled": True,
        "vat_rate": 18,
    }
    p.update(over)
    return p


class TestNewContract:

    def test_happy_path(self, client, created_ids):
        r = client.post(f"{API}/contracts/new", json=_payload())
        assert r.status_code == 200, r.text
        d = r.json()
        assert d["sifarisci_company"] == "TEST_QA MMC"
        assert d["contract_number"] == "TS_QA/27"
        assert d["parent_contract_number"] == "TS_QA/27"
        assert d["stand_m2"] == 12
        assert d["type"] == "new_contract"
        assert "id" in d
        created_ids.append(d["id"])

    def test_docx_swap(self, client, created_ids):
        cid = created_ids[0]
        r = client.get(f"{API}/contracts/{cid}/download")
        assert r.status_code == 200
        assert len(r.content) > 5000
        doc = Document(io.BytesIO(r.content))
        text = "\n".join([p.text for p in doc.paragraphs])
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text += "\n" + cell.text

        found = ["TEST_QA MMC", "9999999999", "TEST Executor", "TS_QA/27",
                 "AZ00TEST0000000000000000TEST", "TEST Bank", "888",
                 "7777777777", "AZ00TEST0000CORR", "TESTAZ2X"]
        missing = ["ANLİFE İNŞAAT", "1009525231", "Haqverdiyev Nazir",
                   "TS002/27", "AZ32TURA40030095266600102944", "TuranBank",
                   "508212", "1300016391",
                   "AZ26NABZ01350100000000027944", "TURAAZ22"]

        for token in found:
            assert token in text, f"expected {token!r} in DOCX text"
        for token in missing:
            assert token not in text, f"unexpected leftover {token!r} in DOCX"

    def test_invoice_xlsx(self, client, created_ids):
        cid = created_ids[0]
        r = client.get(f"{API}/contracts/{cid}/invoice")
        assert r.status_code == 200, r.text
        assert len(r.content) > 3000
        wb = load_workbook(io.BytesIO(r.content))
        assert len(wb.sheetnames) >= 1

    def test_stand_plan(self, client, created_ids):
        cid = created_ids[0]
        r = client.get(f"{API}/contracts/{cid}/stand-plan")
        assert r.status_code == 200, r.text
        assert len(r.content) > 3000

    def test_duplicate_allowed(self, client, created_ids):
        r = client.post(f"{API}/contracts/new",
                        json=_payload(contract_number="TS_QA2/27"))
        assert r.status_code == 200, r.text
        created_ids.append(r.json()["id"])

    def test_missing_company(self, client):
        r = client.post(f"{API}/contracts/new",
                        json=_payload(sifarisci_company=""))
        assert r.status_code == 400
        assert "Şirkət" in r.json().get("detail", "")

    def test_missing_contract_number(self, client):
        r = client.post(f"{API}/contracts/new",
                        json=_payload(contract_number=""))
        assert r.status_code == 400
        assert "Müqavilə nömrəsi" in r.json().get("detail", "")

    def test_addendum_regression(self, client, created_ids):
        r = client.post(f"{API}/contracts/addendum", json={
            "parent_contract_number": "TS_QA/27",
            "parent_contract_date": "2027-02-15",
            "addendum_date": "2027-03-01",
            "sifarisci_company": "TEST_QA MMC",
            "sifarisci_voen": "9999999999",
            "sifarisci_authorized": "TEST Executor",
            "stand_number": "S-99",
            "stand_width": 4,
            "stand_length": 3,
            "exhibition_name": "Yerli sərgi",
            "exhibition_start": "2027-06-23",
            "exhibition_end": "2027-06-26",
            "pricing": {"price_net": 5000, "vat_enabled": True, "vat_rate": 18},
        })
        assert r.status_code == 200, r.text
        created_ids.append(r.json()["id"])
