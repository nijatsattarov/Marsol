"""Contract parsing + addendum generation service.

Handles two kinds of operations:
  1. Extract structured fields from an uploaded "main" service contract (DOCX).
  2. Generate an "addendum" (Müqaviləyə Əlavə) DOCX based on the extracted
     fields plus user-supplied pricing details.

The parser uses heuristic regexes tuned to the Marsol Group contract template
(see /app/memory/PRD.md for canonical samples). Failures are graceful: any
field that can't be detected is returned as an empty string so the user can
fill it in manually in the UI.
"""
from __future__ import annotations

import io
import os
import re
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401  (used indirectly)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage


# ----------------------------------------------------------------------------
# PARSER
# ----------------------------------------------------------------------------

# Patterns are intentionally permissive — Azerbaijani contract numbering can
# look like "№ TS001/26", "№TS001-26", "№ MA-2025/14" etc.
_CONTRACT_NUM_RE = re.compile(r"№\s*([A-Za-zƏəĞğİıÖöÜüÇçŞş0-9\-/_]+)", re.UNICODE)
# Match VÖEN/VOEN with ANY dash-like or colon separator (hyphen, en-dash,
# em-dash, minus, full-width hyphen).
_VOEN_RE = re.compile(r"V[ÖöO]EN[\s:\-–—‒−]*(\d{8,12})", re.IGNORECASE | re.UNICODE)
_DIRECTOR_RE = re.compile(r"direktoru\s+([A-Za-zƏəĞğİıÖöÜüÇçŞş\s\.'-]+?)\s+şəxs", re.UNICODE)
# Contract date — handles DD.MM.YYYY (with optional separators / suffix)
# and `D ay YYYY` style ("14 iyul 2025", "«14» iyul 2025-ci il").
_DATE_RE = re.compile(r"\b(\d{1,2})[.\s/](\d{1,2})[.\s/](\d{4})\b")
_AZ_MONTH_TOKENS = {
    "yanvar": 1, "fevral": 2, "mart": 3, "aprel": 4, "may": 5, "iyun": 6,
    "iyul": 7, "avqust": 8, "sentyabr": 9, "oktyabr": 10, "noyabr": 11, "dekabr": 12,
}
_AZ_DATE_RE = re.compile(
    r"[«\"']?(\d{1,2})[»\"']?\s+("
    + "|".join(_AZ_MONTH_TOKENS.keys())
    + r")\s+(\d{4})", re.IGNORECASE
)
# Company names usually appear wrapped in “ ” / " " / « » with the legal-form
# suffix (MMC, QSC, ASC, "Məhdud Məsuliyyətli Cəmiyyəti", ...). We grab the
# bit inside the guillemets.
_COMPANY_RE = re.compile(
    r"[“\"„«]\s*([A-Za-zƏəĞğİıÖöÜüÇçŞş0-9\s\.&\-]+?)\s*[”\"“»]\s*"
    r"(?:QSC|MMC|ASC|Məhdud|Qapalı|Açıq|Limited)",
    re.UNICODE | re.IGNORECASE
)


def _read_docx_text(file_bytes: bytes) -> str:
    """Return all paragraphs + table cells joined with newlines."""
    doc = Document(io.BytesIO(file_bytes))
    chunks: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    chunks.append(t)
    return "\n".join(chunks)


def extract_contract_fields(file_bytes: bytes) -> Dict[str, str]:
    """Best-effort extraction. Returns dict with keys:
      contract_number, sifarisci_company, sifarisci_voen, sifarisci_authorized,
      icraci_company, icraci_voen, icraci_authorized, raw_text
    Any missing field comes back as empty string."""
    text = _read_docx_text(file_bytes)

    contract_number = ""
    m = _CONTRACT_NUM_RE.search(text)
    if m:
        contract_number = m.group(1).strip()

    # Contract date — pick the FIRST DD.MM.YYYY occurrence (header date).
    # Returned in ISO format (YYYY-MM-DD) so it slots straight into a
    # `<input type="date">` on the frontend.
    contract_date_iso = ""
    # Skip dates that look like placeholders ("01.01.0000", etc.) by requiring
    # a recent year (>= 2015) which all real Marsol contracts have.
    for dm in _DATE_RE.finditer(text):
        try:
            dd, mm, yy = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
            if 1 <= dd <= 31 and 1 <= mm <= 12 and 2015 <= yy <= 2099:
                contract_date_iso = f"{yy:04d}-{mm:02d}-{dd:02d}"
                break
        except ValueError:
            continue
    # Fallback: try `14 iyul 2025` Azerbaijani-month form
    if not contract_date_iso:
        adm = _AZ_DATE_RE.search(text)
        if adm:
            try:
                dd = int(adm.group(1))
                mm = _AZ_MONTH_TOKENS[adm.group(2).lower()]
                yy = int(adm.group(3))
                if 2015 <= yy <= 2099:
                    contract_date_iso = f"{yy:04d}-{mm:02d}-{dd:02d}"
            except (ValueError, KeyError):
                pass

    # Find ALL VÖEN matches (usually 2: Sifarişçi + İcraçı). Marsol's VÖEN is
    # known to be 2004204701, so we use it to determine which side is which.
    voens = _VOEN_RE.findall(text)
    voens = list(dict.fromkeys(voens))  # preserve order, dedupe
    marsol_voen = "2004204701"
    sifarisci_voen = ""
    icraci_voen = marsol_voen
    for v in voens:
        if v != marsol_voen and not sifarisci_voen:
            sifarisci_voen = v

    # Director names — usually 2 of them
    directors = _DIRECTOR_RE.findall(text)
    directors = [d.strip() for d in directors]
    # Heuristic: director that appears with Marsol's VÖEN goes to icraci.
    # If we can't tell, the FIRST director in the doc is sifarisci.
    sifarisci_authorized = ""
    icraci_authorized = "Bilal Qasımlı"  # Marsol default
    for d in directors:
        if "qasım" not in d.lower() and "marsol" not in d.lower() and not sifarisci_authorized:
            sifarisci_authorized = d

    # Company names — pick first 2 distinct
    companies = _COMPANY_RE.findall(text)
    companies = [c.strip() for c in companies]
    companies = list(dict.fromkeys(companies))
    sifarisci_company = ""
    icraci_company = "MARSOL"
    for c in companies:
        if "marsol" not in c.lower() and not sifarisci_company:
            sifarisci_company = c

    return {
        "contract_number": contract_number,
        "contract_date": contract_date_iso,
        "sifarisci_company": sifarisci_company,
        "sifarisci_voen": sifarisci_voen,
        "sifarisci_authorized": sifarisci_authorized,
        "icraci_company": icraci_company,
        "icraci_voen": icraci_voen,
        "icraci_authorized": icraci_authorized,
        "raw_text_preview": text[:1500],
    }


# ----------------------------------------------------------------------------
# GENERATOR — Addendum (Müqaviləyə Əlavə)
# ----------------------------------------------------------------------------

# Marsol İcraçı reqviziləri default-da bunlardır. Üzərinə yazmaq mümkündür.
MARSOL_DEFAULT = {
    "name": "MARSOL",
    "full_name": "MARSOL Məhdud Məsuliyyətli Cəmiyyəti",
    "voen": "2004204701",
    "iban": "AZ41AIIB40060019440152635107",
    "bank_name": "Kapital Bank ASC Xətai filialı",
    "bank_branch_code": "200071",
    "bank_voen": "9900003611",
    "corr_account": "AZ37NABZ01350100000000001944",
    "swift": "AIIBAZ2X",
    "authorized": "Bilal Qasımlı",
}


# Standart sərgi paketinə daxil olan xidmətlər. Bu siyahı statikdir və
# istifadəçi tərəfindən UI-də dəyişdirilə bilməz — Marsol Group-un standart
# sərgi iştirakı paketinə uyğun olaraq müqavilədə həmişə eyni cədvəl çıxır.
STANDARD_SERVICES: List[str] = [
    "Arxa və yan divarlar, şirkətin adı ilə lövhə, 1 masa, 2 stul, zibil qutusu, 3 yuvalı elektrik uzadıcı, xalça döşəmə;",
    "Sərgi üçün çap ediləcək 2000 tiraj kataloqda 1(bir) səhifə reklam (A5).",
    "Sertifikat",
    "Sərgi foye hissəsində 6x4 metr ölçülü monitorda vaxtaşırı şirkətinizin loqosunun yayımlanması",
    "“Brendwall” da Sifarişçiyə məxsus loqo",
    "Marsolexpo.az saytında şirkətiniz üçün ayrılmış bölmədə məlumatlarınızın bir illik yerləşdirilməsi",
    "“Sərgidə biz də varıq” posterinin tərəfimizdən tərtib olunması;",
    "Sərgi günlərində təşkil olunan B2B və B2G görüşlərdə iştirak imkanı;",
    "Coffee Break zonasında təqdim olunan xidmətlərdən ödənişsiz istifadə (ancaq stend iştirakçıları üçün)",
    "Sərgi sonrası axşam ziyafətinə bir nəfərə dəvətnamə",
    "Axşam ziyafətində “Brendwall” da loqo",
    "Sərgi müddətində, vaxtaşırı kampaniya və endirimlərin səsləndirilməsi;",
    "Sərgi iştirakçısı şirkət rəhbərlərilə sərgi öncəsi təşkil edilən görüşlərdə iştirak imkanı;",
]


def _fmt_date_az(iso_date: str) -> str:
    """`2025-07-14` → `14.07.2025`. Empty / invalid → '__.__.____'."""
    if not iso_date:
        return "__.__.____"
    try:
        dt = datetime.strptime(iso_date[:10], "%Y-%m-%d")
        return dt.strftime("%d.%m.%Y")
    except (ValueError, TypeError):
        return iso_date


def _fmt_money(value) -> str:
    """`12345.6` → `12 345.00` (AZ-style space thousand separator)."""
    try:
        n = float(value)
    except (TypeError, ValueError):
        return ""
    return f"{n:,.2f}".replace(",", " ")


# --- Azerbaijani number → words (for "Məbləğ yazı ilə") ---------------------
_AZ_ONES = ["", "bir", "iki", "üç", "dörd", "beş", "altı", "yeddi", "səkkiz", "doqquz"]
_AZ_TENS = ["", "on", "iyirmi", "otuz", "qırx", "əlli", "altmış", "yetmiş",
            "səksən", "doxsan"]


def _az_int_to_words(n: int) -> str:
    """Convert an integer (0..999 999 999) to Azerbaijani words."""
    if n == 0:
        return "sıfır"
    if n < 0:
        return "mənfi " + _az_int_to_words(-n)

    def _under_thousand(x: int) -> str:
        parts = []
        h, rem = divmod(x, 100)
        if h:
            parts.append("yüz" if h == 1 else f"{_AZ_ONES[h]} yüz")
        t, o = divmod(rem, 10)
        if t:
            parts.append(_AZ_TENS[t])
        if o:
            parts.append(_AZ_ONES[o])
        return " ".join(parts)

    parts = []
    millions, rem = divmod(n, 1_000_000)
    if millions:
        parts.append(_under_thousand(millions) + " milyon")
    thousands, rem = divmod(rem, 1000)
    if thousands:
        parts.append(("min" if thousands == 1 else _under_thousand(thousands) + " min"))
    if rem:
        parts.append(_under_thousand(rem))
    return " ".join(parts).strip()


def _az_money_in_words(amount) -> str:
    """`9440.00` → `doqquz min dörd yüz qırx manat 00 qəpik`."""
    try:
        a = float(amount)
    except (TypeError, ValueError):
        return ""
    if a <= 0:
        return "sıfır manat 00 qəpik"
    manat = int(a)
    qepik = int(round((a - manat) * 100))
    return f"{_az_int_to_words(manat)} manat {qepik:02d} qəpik"


_AZ_MONTHS = {
    1: "yanvar", 2: "fevral", 3: "mart", 4: "aprel", 5: "may", 6: "iyun",
    7: "iyul", 8: "avqust", 9: "sentyabr", 10: "oktyabr", 11: "noyabr", 12: "dekabr",
}


def _az_year_suffix(year: int) -> str:
    """Azerbaijani ordinal suffix for a year, e.g. 2025 -> 'ci', 2026 -> 'cı'."""
    return {0: "cu", 1: "ci", 2: "ci", 3: "cü", 4: "cü",
            5: "ci", 6: "cı", 7: "ci", 8: "ci", 9: "cu"}.get(year % 10, "cı")


def _exhibition_range_az(start_iso: str, end_iso: str) -> str:
    """Return 'DD month - DD month YYYY-cI il' style range."""
    if not start_iso or not end_iso:
        return "__ ____ - __ ____ 202_-cı il"
    try:
        s = datetime.strptime(start_iso[:10], "%Y-%m-%d")
        e = datetime.strptime(end_iso[:10], "%Y-%m-%d")
    except ValueError:
        return f"{start_iso} - {end_iso}"
    suffix = _az_year_suffix(e.year)
    if s.year == e.year and s.month == e.month:
        return f"{s.day} - {e.day} {_AZ_MONTHS[e.month]} {e.year}-{suffix} il"
    if s.year == e.year:
        return f"{s.day} {_AZ_MONTHS[s.month]} - {e.day} {_AZ_MONTHS[e.month]} {e.year}-{suffix} il"
    return f"{s.day} {_AZ_MONTHS[s.month]} {s.year} - {e.day} {_AZ_MONTHS[e.month]} {e.year}-{suffix} il"


def _replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    """In-place replace text in a paragraph while preserving run formatting.

    Strategy: iterate over runs. For each placeholder, find the first run that
    contains it and rewrite that run's text. Handles single-run placeholders
    (which is the case for our template). For cross-run placeholders we fall
    back to concatenating + writing to first run (loses inner formatting but
    keeps the paragraph style).
    """
    for old, new in replacements.items():
        # Fast path: any single run contains the placeholder
        hit = False
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                hit = True
                break  # only replace first occurrence per call
        if hit:
            continue
        # Slow path: placeholder split across runs — flatten paragraph text
        full = "".join(r.text for r in paragraph.runs)
        if old in full:
            new_full = full.replace(old, new, 1)
            if paragraph.runs:
                paragraph.runs[0].text = new_full
                for r in paragraph.runs[1:]:
                    r.text = ""


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 12) -> None:
    """Replace cell content with a single Times New Roman run."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def generate_addendum_docx(data: Dict) -> bytes:
    """Generate the addendum DOCX by filling the Marsol template.

    Only the empty placeholder cells (parent contract no/date, addendum no/date,
    Sifarişçi company / VÖEN / authorized person, exhibition dates) are
    replaced. ALL other text, fonts, tables, services list and signature block
    layout come straight from the user-supplied template at
    `/app/backend/templates/addendum_template.docx`.
    """
    import os

    template_path = os.path.join(os.path.dirname(__file__), "templates",
                                 "addendum_template.docx")
    doc = Document(template_path)

    # --- Build replacement values ---
    addendum_no = (data.get("addendum_number") or "1").strip()
    addendum_date = _fmt_date_az(data.get("addendum_date")
                                 or datetime.now().strftime("%Y-%m-%d"))
    parent_no = (data.get("parent_contract_number") or "____").strip()
    parent_date = _fmt_date_az(data.get("parent_contract_date"))
    sif_co = (data.get("sifarisci_company") or "").strip()
    sif_voen = (data.get("sifarisci_voen") or "").strip()
    sif_auth = (data.get("sifarisci_authorized") or "").strip()

    ex_range = _exhibition_range_az(data.get("exhibition_start"),
                                    data.get("exhibition_end"))
    # If user didn't supply exhibition dates, keep the static template text
    # as-is (don't overwrite "23 iyun - 26 iyun 2027-ci il" with underscores).
    skip_exhibition = not (data.get("exhibition_start") and data.get("exhibition_end"))

    # Suffix for parent contract date (e.g. "-cı il")
    p_year = None
    try:
        p_year = datetime.strptime((data.get("parent_contract_date") or "")[:10],
                                   "%Y-%m-%d").year
    except (ValueError, TypeError):
        p_year = None
    p_year_suffix = _az_year_suffix(p_year) if p_year else "cı"
    parent_date_full = (f"{parent_date}-{p_year_suffix} il"
                        if p_year else "__.___.202_-cı il")
    parent_num_full = f"{parent_no} №-li" if parent_no and parent_no != "____" \
        else "_____ №-li"

    # Suffix for addendum date
    a_year = None
    try:
        a_year = datetime.strptime((data.get("addendum_date") or "")[:10],
                                   "%Y-%m-%d").year
    except (ValueError, TypeError):
        a_year = None
    a_year_suffix = _az_year_suffix(a_year) if a_year else "cı"
    addendum_date_full = (f"{addendum_date}-{a_year_suffix} il"
                          if a_year else "__.___. 202_-cı il")

    # Sifarişçi şirkət adı: template-də placeholder-dan sonra "Məhdud
    # Məsuliyyətli Cəmiyyəti" sabit yazılır. Əgər istifadəçinin daxil etdiyi
    # şirkət adı artıq hüquqi forma daşıyırsa (QSC, ASC, MMC, MMC-si və s.),
    # template-də sonradan gələn " Məhdud Məsuliyyətli Cəmiyyəti" ifadəsini
    # də siləcəyik. Bunu daha geniş placeholder ilə birdəfəlik əvəz edirik.
    has_legal_form = bool(re.search(
        r"\b(MMC|QSC|ASC|HM|FH|ÖC|MM[CV]|Məhdud|Açıq|Qapalı)\b",
        sif_co or "", re.IGNORECASE
    ))
    sif_company_full = (
        f"«{sif_co}»" if has_legal_form and sif_co
        else (f"«{sif_co}» Məhdud Məsuliyyətli Cəmiyyəti" if sif_co
              else "“__________ ” Məhdud Məsuliyyətli Cəmiyyəti")
    )

    # The template uses these EXACT placeholder strings (copied from the
    # final version sent by the user). They must match byte-for-byte.
    replacements = {
        # Parent contract reference – appears multiple times
        "__.___.202_-cı il": parent_date_full,
        "_____ №-li": parent_num_full,
        # Addendum number (header)
        "ƏLAVƏ №___": f"ƏLAVƏ №{addendum_no}",
        # Addendum date (in P3, note the extra space before "202_")
        "__.___. 202_-cı il": addendum_date_full,
        # Sifarişçi company name (P9 — the full phrase including the legal
        # form so that companies like QSC/ASC don't get a duplicate suffix).
        # NOTE: template has TWO spaces between `”` and `Məhdud`.
        "“__________ ”  Məhdud Məsuliyyətli Cəmiyyəti": sif_company_full,
        # Sifarişçi VÖEN (P10)
        "VÖEN – ____________": f"VÖEN – {sif_voen}" if sif_voen else "VÖEN – ____________",
        # Sifarişçi authorized person (P11)
        "səlahiyyətli şəxs – __________": (
            f"səlahiyyətli şəxs – {sif_auth}" if sif_auth
            else "səlahiyyətli şəxs – __________"
        ),
        # Exhibition date range (P15) — only replaced when user supplied dates
        **({} if skip_exhibition else {"23 iyun - 26 iyun 2027-ci il": ex_range}),
    }

    # --- Replace inside paragraphs ---
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    # --- Replace inside table cells (some placeholders may live in tables) ---
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replacements)

    # --- Fill Sifarişçi column in the signature table (Table 1) ---
    if len(doc.tables) >= 2:
        sig_tbl = doc.tables[1]
        # Expected layout (3 cols): [label, Sifarişçi value, İcraçı value]
        # R0: header, R1: company, R2: VÖEN, R3: səlahiyyətli şəxs,
        # R4: İmza, R5: Möhür
        try:
            if sif_co:
                _set_cell_text(sig_tbl.rows[1].cells[1],
                               f"«{sif_co}»", bold=True)
            if sif_voen:
                _set_cell_text(sig_tbl.rows[2].cells[1], sif_voen)
            if sif_auth:
                _set_cell_text(sig_tbl.rows[3].cells[1], sif_auth)
        except (IndexError, KeyError):
            pass

    # --- Insert pricing table + "Məbləğ yazı ilə" line right AFTER the
    # services table. The template that the user provided does not contain
    # these elements, so we synthesize them every time. Only ONE pricing
    # table is added (no duplicates).
    pricing = data.get("pricing", {}) or {}
    price_net = float(pricing.get("price_net") or 0)
    vat_enabled = bool(pricing.get("vat_enabled", True))
    vat_rate = float(pricing.get("vat_rate") or 18) if vat_enabled else 0
    vat_amount = round(price_net * vat_rate / 100, 2) if vat_enabled else 0.0
    total = round(price_net + vat_amount, 2)
    stand_no = (data.get("stand_number") or "").strip()
    service_label = (data.get("service_label") or "Sərgidə iştirak").strip()
    unit_label = (data.get("unit_label") or "x/h").strip()

    if len(doc.tables) >= 1:
        services_tbl = doc.tables[0]

        # 1) Pricing table — 2 rows × 6 cols, gray header row
        price_tbl = doc.add_table(rows=2, cols=6)
        price_tbl.style = "Table Grid"
        headers = ["Stend №", "Xidmətin adı", "Ölçü vahidi",
                   "Qiymət", "ƏDV", "Yekun məbləğ"]
        for i, h in enumerate(headers):
            cell = price_tbl.rows[0].cells[i]
            _set_cell_text(cell, h, bold=True, size=11)
            # Apply gray shading (D9D9D9) to header cell
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "D9D9D9")
            tc_pr.append(shd)

        data_row = [
            stand_no or "",
            service_label,
            unit_label,
            _fmt_money(price_net) if price_net else "",
            (_fmt_money(vat_amount) if vat_enabled and vat_amount else "-"),
            _fmt_money(total) if total else "-",
        ]
        for i, v in enumerate(data_row):
            _set_cell_text(price_tbl.rows[1].cells[i], v, size=11)

        # Move price table to right after services table in document order
        svc_xml = services_tbl._element
        price_xml = price_tbl._element
        price_xml.getparent().remove(price_xml)
        svc_xml.addnext(price_xml)

        # 2) "Məbləğ yazı ilə:" paragraph immediately after the price table
        money_para = OxmlElement("w:p")
        # Paragraph properties (left-aligned, normal spacing)
        p_pr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "120")
        spacing.set(qn("w:after"), "120")
        p_pr.append(spacing)
        money_para.append(p_pr)

        def _add_run(parent, text: str, *, bold: bool = False):
            run = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                r_fonts.set(qn(attr), "Times New Roman")
            r_pr.append(r_fonts)
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); r_pr.append(sz)
            szcs = OxmlElement("w:szCs"); szcs.set(qn("w:val"), "24"); r_pr.append(szcs)
            if bold:
                b = OxmlElement("w:b"); r_pr.append(b)
                bcs = OxmlElement("w:bCs"); r_pr.append(bcs)
            run.append(r_pr)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = text
            run.append(t)
            parent.append(run)

        in_words = _az_money_in_words(total) if total else "sıfır manat 00 qəpik"
        _add_run(money_para, "Məbləğ yazı ilə: ", bold=True)
        _add_run(money_para, in_words)

        price_xml.addnext(money_para)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ----------------------------------------------------------------------------
# INVOICE (HESAB-FAKTURA) — XLSX GENERATION
# ----------------------------------------------------------------------------

MARSOL_HEADER_PATH = os.path.join(
    os.path.dirname(__file__), "templates", "marsol_header.png"
)

INVOICE_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "templates", "invoice_template.xlsx"
)


def _apply_marsol_header(ws) -> None:
    """Stamp the Marsol branding header onto the top of the invoice sheet.

    We embed the user-supplied `marsol_header.png` as a single image
    anchored at A1. This avoids any text/logo overlap because the entire
    branding (logo + address + phones + web/email + colored divider) is
    pre-rendered in the image. Rows 1-6 are reserved (row heights raised)
    so the image doesn't visually collide with the HESAB-FAKTURA title at
    row 8.
    """
    if not os.path.exists(MARSOL_HEADER_PATH):
        return
    try:
        img = XLImage(MARSOL_HEADER_PATH)
        # Source image is a wide banner (~1315×240, ratio ≈ 5.48). Scale to
        # ~700 px wide so it spans the full A..G width without distortion.
        img.width = 720
        img.height = 132
        img.anchor = "A1"
        ws.add_image(img)
    except (FileNotFoundError, OSError):
        return

    # Reserve enough vertical space (≈ 132 px ≈ 7 rows × 22 px). We use
    # rows 1-5 here so the HESAB-FAKTURA title at row 8 stays untouched.
    for r in range(1, 6):
        ws.row_dimensions[r].height = 28


def generate_invoice_xlsx(data: Dict) -> bytes:
    """Build a Hesab-Faktura XLSX based on the same data used for the addendum.

    The template at `INVOICE_TEMPLATE_PATH` is the user-provided HF TS058.xlsx
    file with two sheets ('sərgi' = main invoice, 'reqem' = number-to-words
    helper sheet with cascading formulas). We ONLY overwrite the customer-
    specific cells; all formulas, styles, layout and bank details stay
    untouched.

    Cell map (sheet 'sərgi'):
      A8  — `HESAB-FAKTURA No <müqavilə №>`
      A15 — ` Hesabı alan:<Sifarişçi adı>`
      A16 — `VÖEN: <Sifarişçi VÖEN>`
      G13 — invoice date (Tarix)         → today / addendum_date
      G14 — Müqavilə №                    → parent_contract_number
      G15 — Müqavilənin tarixi            → parent_contract_date
      E20 — Miqdar                        → 1
      F20 — Qiymət (ƏDV-siz)              → pricing.price_net
      (Cəmi / ƏDV / Cəmi (ƏDV-ilə) / sözlərlə cəmi — formulas auto-compute.)
    """
    wb = load_workbook(INVOICE_TEMPLATE_PATH)
    ws = wb["sərgi"] if "sərgi" in wb.sheetnames else wb.active

    # --- Header (Marsol logo + contact bar) ---
    _apply_marsol_header(ws)

    parent_no = (data.get("parent_contract_number") or "").strip() or "____"
    sif_co = (data.get("sifarisci_company") or "").strip()
    sif_voen = (data.get("sifarisci_voen") or "").strip()

    pricing = data.get("pricing", {}) or {}
    price_net = float(pricing.get("price_net") or 0)

    # Parse ISO dates into Python datetime so Excel keeps native cell type.
    def _to_dt(iso: Optional[str]) -> Optional[datetime]:
        if not iso:
            return None
        try:
            return datetime.strptime(iso[:10], "%Y-%m-%d")
        except (ValueError, TypeError):
            return None

    invoice_date = _to_dt(data.get("addendum_date")) or datetime.now()
    parent_date = _to_dt(data.get("parent_contract_date"))

    # 1) Invoice header — keep original spacing/format
    ws["A8"] = f"HESAB-FAKTURA No {parent_no} "

    # 2) Sifarişçi (customer) block
    #    Original sample had a leading space in A15 — keep it for layout parity.
    if sif_co:
        ws["A15"] = f" Hesabı alan: {sif_co}"
    if sif_voen:
        ws["A16"] = f"VÖEN: {sif_voen}"

    # 3) Right-side meta block
    ws["G13"] = invoice_date
    ws["G14"] = parent_no
    if parent_date:
        ws["G15"] = parent_date

    # 4) Services row (only the variable bits — the rest is hardcoded)
    ws["E20"] = 1                  # Miqdar
    if price_net:
        ws["F20"] = price_net      # Qiymət (ƏDV-siz)

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ----------------------------------------------------------------------------
# STAND PLAN (STEND YERLƏŞİM PLANI) — DOCX GENERATION
# ----------------------------------------------------------------------------

STAND_PLAN_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "templates", "stand_plan_template.docx"
)


def generate_stand_plan_docx(data: Dict) -> bytes:
    """Build the stand-placement plan DOCX based on the addendum data.

    The template at `STAND_PLAN_TEMPLATE_PATH` contains:
      • P0 — Title (static)
      • P3 — `Stend № <N>  en –<W> m ; uzunluq – <L> m ;  ümumi sahə - <A> m²`
      • Table R0 C0 — `"<Company>"MMC / Sahibkar: <Owner Name>     M.`
      • Table R0 C1 — Marsol / Direktor: Bilal Qasımlı (static)

    We replace the customer-specific placeholders while keeping all
    formatting (fonts, table, alignment, image, etc.) intact.
    """
    doc = Document(STAND_PLAN_TEMPLATE_PATH)

    sif_co = (data.get("sifarisci_company") or "").strip()
    sif_owner = (data.get("sifarisci_authorized") or "").strip()
    stand_no = (data.get("stand_number") or "").strip()

    # Numeric fields (stand width/length in metres) — area = w × l
    try:
        width = float(data.get("stand_width") or 0)
    except (TypeError, ValueError):
        width = 0.0
    try:
        length = float(data.get("stand_length") or 0)
    except (TypeError, ValueError):
        length = 0.0
    area = round(width * length, 2)

    # Format numbers without trailing ".0" if integer
    def _num(x: float) -> str:
        return f"{int(x)}" if x and float(x).is_integer() else f"{x:g}"

    # --- Replace P3 (the line with stand number / dimensions / area) ---
    # Old text: "Stend № 70  en –6 m ; uzunluq – 8 m ;  ümumi sahə - 48 m² "
    new_p3 = (
        f"Stend № {stand_no or '___'}  "
        f"en –{_num(width) if width else '___'} m ; "
        f"uzunluq – {_num(length) if length else '___'} m ; "
        f" ümumi sahə - {_num(area) if area else '___'} m² "
    )
    for p in doc.paragraphs:
        if "Stend №" in p.text and ("en –" in p.text or "ümumi sahə" in p.text):
            if p.runs:
                p.runs[0].text = new_p3
                for r in p.runs[1:]:
                    r.text = ""
            break

    # --- Replace Sifarişçi cell (R0 C0 in Table 0) ---
    if doc.tables:
        cell = doc.tables[0].rows[0].cells[0]
        full = "\n".join(p.text for p in cell.paragraphs)
        # Original: `“PROMO EXPO”MMC / Sahibkar: Araz Mahmudov      M.`
        # New:      `“<COMPANY>”MMC / Sahibkar: <OWNER>             M.`
        new_text = f"“{sif_co or '__________'}” / Sahibkar: {sif_owner or '__________'}     M."
        # Replace text in the first paragraph + first run while keeping format
        first_p = cell.paragraphs[0]
        if first_p.runs:
            first_p.runs[0].text = new_text
            for r in first_p.runs[1:]:
                r.text = ""
        else:
            first_p.text = new_text
        # Clear remaining paragraphs (template originally only had 1)
        for extra_p in cell.paragraphs[1:]:
            if extra_p.runs:
                for r in extra_p.runs:
                    r.text = ""
        del full  # noqa

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()